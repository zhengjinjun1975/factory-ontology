#!/usr/bin/env python3
"""knowledge/ingest.py — 文档解析（PDF/Word/TXT → 纯文本）

把甲方交付的文档（.pdf/.docx/.txt）解析为 {title, raw_text}。
PDF 优先 pymupdf(fitz)，缺失则试 pdfplumber；docx 用 python-docx；txt 直接读。
任一依赖缺失或解析失败都返回 None（不抛异常），由调用方决定是否降级。

用法:
    from knowledge.ingest import extract_text
    doc = extract_text("甲方技术协议.pdf")   # {title, raw_text} 或 None
"""
import os


def extract_text(path):
    """解析文档为 {title, raw_text}。缺库/解析失败返回 None（不抛异常）。

    title 取文件名（去扩展名）；raw_text 为解析出的纯文本。
    """
    try:
        if not path or not os.path.isfile(path):
            return None
        title = os.path.splitext(os.path.basename(path))[0]
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            raw = _extract_pdf(path)
        elif ext in (".docx", ".doc"):
            raw = _extract_docx(path)
        elif ext == ".txt":
            raw = _extract_txt(path)
        else:
            raw = None
        if not raw or not raw.strip():
            return None
        return {"title": title, "raw_text": raw.strip()}
    except Exception:
        return None


_rapid_ocr = None  # RapidOCR 实例缓存（惰性初始化，复用避免每次重载模型）


def _get_rapid_ocr():
    """惰性获取 RapidOCR 单例。缺库/初始化失败返回 None（不抛异常）。"""
    global _rapid_ocr
    if _rapid_ocr is not None:
        return _rapid_ocr
    try:
        from rapidocr_onnxruntime import RapidOCR
        _rapid_ocr = RapidOCR()  # 首次初始化，之后复用
        return _rapid_ocr
    except Exception:
        return None


def _winsdk_ocr(png_path):
    """Windows 自带 OCR(winsdk, 免装模型) 识别中文。返回文本或 None。"""
    import asyncio
    try:
        from winsdk.windows.media.ocr import OcrEngine
        from winsdk.windows.graphics.imaging import BitmapDecoder
        from winsdk.windows.storage import StorageFile, FileAccessMode
        from winsdk.windows.globalization import Language

        async def _ocr():
            file = await StorageFile.get_file_from_path_async(png_path)
            stream = await file.open_async(FileAccessMode.READ)
            decoder = await BitmapDecoder.create_async(stream)
            bitmap = await decoder.get_software_bitmap_async()
            engine = OcrEngine.try_create_from_language(Language("zh-CN"))
            if engine is None:
                engine = OcrEngine.try_create_from_user_profile_languages()
            if engine is None:
                return None
            result = await engine.recognize_async(bitmap)
            return result.text.strip() if result and result.text else None
        return asyncio.run(_ocr())
    except Exception:
        return None


def _ocr_page_image(page, dpi=150):
    """把 PDF 页渲染为图片，用 Windows 自带 OCR(winsdk, 免装模型) 识别中文；
    winsdk 不可用时回退 RapidOCR。失败返回 None。"""
    import os, tempfile
    img_bytes = None
    try:
        pix = page.get_pixmap(dpi=dpi)
        img_bytes = pix.tobytes("png")
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        try:
            tmp.write(img_bytes)
            tmp.close()
            text = _winsdk_ocr(tmp.name)
            if text:
                return text
        finally:
            try:
                os.unlink(tmp.name)
            except OSError:
                pass
    except Exception:
        pass
    # 回退 RapidOCR
    try:
        ocr = _get_rapid_ocr()
        if ocr is None or img_bytes is None:
            return None
        result, _elapsed = ocr(img_bytes)
        if not result:
            return None
        lines = []
        for item in result or []:
            if isinstance(item, (list, tuple)) and len(item) >= 2:
                txt = item[1]
                if isinstance(txt, (list, tuple)):
                    txt = txt[0] if txt else ""
                if txt:
                    lines.append(str(txt))
        return "\n".join(lines)
    except Exception:
        return None


def _extract_pdf(path):
    """PDF → 文本。优先 pymupdf(fitz)，其次 pdfplumber。

    若某页文本层为空（扫描版/图片型 PDF），自动用 RapidOCR 识别该页图片，
    文本层 + OCR 结果合并为该页内容。OCR 缺失或失败时优雅降级回原逻辑。
    """
    try:
        import fitz  # pymupdf
        text_parts = []
        with fitz.open(path) as doc:
            for page in doc:
                layer_text = page.get_text()  # 文本层
                if layer_text and layer_text.strip():
                    # 有文本层：直接用，不 OCR
                    text_parts.append(layer_text)
                else:
                    # 扫描页：文本层为空，走 OCR 识别图片
                    ocr_text = _ocr_page_image(page)
                    if ocr_text:
                        text_parts.append(ocr_text)
        return "\n".join(p for p in text_parts if p and p.strip())
    except ImportError:
        pass
    except Exception:
        return None
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join((p.extract_text() or "") for p in pdf.pages)
    except ImportError:
        return None
    except Exception:
        return None


def _extract_docx(path):
    """Word .docx → 文本（python-docx）。"""
    try:
        import docx
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
        for t in d.tables:
            for row in t.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception:
        return None


def _extract_txt(path):
    """TXT → 文本。兼容 UTF-8 / GBK / 常见编码。"""
    for enc in ("utf-8", "gbk", "gb18030", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception:
            return None
    return None
